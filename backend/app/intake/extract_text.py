"""Document text extraction — the first pipeline stage, always local.

Supported inputs:
  * .txt / .md / .rtf  -> decoded text (RTF control words stripped)
  * .pdf               -> embedded text via PyMuPDF; if a page has no text layer
                          (scanned image), OCR is attempted per page
  * .docx              -> paragraphs via python-docx
  * image (.png/.jpg)  -> OCR

OCR is OPTIONAL and pluggable. If Tesseract/pytesseract is not installed, scanned
PDFs and images degrade gracefully: we return whatever text exists plus a clear
warning, rather than failing. All processing happens on this machine; nothing is
uploaded for text extraction.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field


@dataclass
class ExtractedDocument:
    text: str
    source_kind: str          # "text" | "pdf" | "pdf+ocr" | "docx" | "image" | "fhir" | "ccda"
    warnings: list[str] = field(default_factory=list)
    ocr_used: bool = False


_RTF_CONTROL = re.compile(r"\\[a-zA-Z]+-?\d* ?")
_RTF_HEX = re.compile(r"\\'([0-9a-fA-F]{2})")

# Accepted upload types and their magic-byte signatures. Text-ish formats have no
# reliable signature (validated as decodable instead).
ALLOWED_EXTENSIONS = (
    ".txt", ".md", ".rtf", ".pdf", ".docx", ".json", ".fhir", ".ndjson", ".xml",
    ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp",
)
_MAGIC = {
    ".pdf": (b"%PDF",),
    ".docx": (b"PK\x03\x04",),          # docx is a zip
    ".png": (b"\x89PNG\r\n\x1a\n",),
    ".jpg": (b"\xff\xd8\xff",), ".jpeg": (b"\xff\xd8\xff",),
    ".tif": (b"II*\x00", b"MM\x00*"), ".tiff": (b"II*\x00", b"MM\x00*"),
    ".bmp": (b"BM",),
}
# Byte-order marks some exporters prepend; stripped before signature matching so a
# BOM-prefixed file isn't misread as "wrong content".
_BOMS = (b"\xef\xbb\xbf", b"\xff\xfe", b"\xfe\xff")
# PDFs are the one format where the marker legitimately floats: Adobe and PyMuPDF
# both accept "%PDF" anywhere in the first kilobyte (leading whitespace/junk from a
# scanner or exporter is common). Match the parser's tolerance so validation never
# rejects a file the extractor could actually read.
_SIGNATURE_SCAN_BYTES = 1024


def _strip_bom(data: bytes) -> bytes:
    for bom in _BOMS:
        if data.startswith(bom):
            return data[len(bom):]
    return data


class UploadRejected(ValueError):
    """Raised when an upload fails validation. Carries a client-safe message + code."""

    def __init__(self, message: str, status_code: int = 415) -> None:
        super().__init__(message)
        self.status_code = status_code


def validate_upload(filename: str, data: bytes) -> None:
    """Fail closed on empty, unsupported, or content-type-mismatched uploads. Message
    text is client-safe (no chart content)."""
    name = (filename or "").lower()
    if not data:
        raise UploadRejected("Empty file.", 400)
    ext = "." + name.rsplit(".", 1)[-1] if "." in name else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise UploadRejected(f"Unsupported file type {ext or '(none)'}. Allowed: "
                             + ", ".join(sorted(set(ALLOWED_EXTENSIONS))), 415)
    sigs = _MAGIC.get(ext)
    if sigs:
        head = _strip_bom(data)
        # PDFs: scan the first KB (marker can float). Other binary formats really do
        # start with their signature, so require it at the front.
        if ext == ".pdf":
            matched = any(sig in head[:_SIGNATURE_SCAN_BYTES] for sig in sigs)
        else:
            matched = any(head.startswith(sig) for sig in sigs)
        if not matched:
            kind = ext.lstrip(".").upper()
            raise UploadRejected(
                f"This file is named \"{ext}\" but its contents aren't a valid {kind} "
                f"(it may be renamed, corrupted, or actually a different format). "
                f"Re-save it as a real {kind}, or paste the report text into the box below.",
                415,
            )
    if ext == ".pdf" and b"/Encrypt" in data[:4096]:
        raise UploadRejected(
            "This PDF is password-protected, so its text can't be read. "
            "Provide an unlocked copy, or paste the report text into the box below.",
            415,
        )


def extract_text(filename: str, data: bytes) -> ExtractedDocument:
    name = (filename or "").lower()
    if name.endswith((".txt", ".md")):
        return ExtractedDocument(text=_decode(data), source_kind="text")
    if name.endswith(".rtf"):
        return ExtractedDocument(text=_strip_rtf(_decode(data)), source_kind="text")
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith((".json", ".fhir", ".ndjson")):
        from app.intake.fhir_ingest import render_fhir
        text, warnings = render_fhir(data)
        return ExtractedDocument(text=text, source_kind="fhir", warnings=warnings)
    if name.endswith(".xml"):
        from app.intake.fhir_ingest import render_ccda
        text, warnings = render_ccda(data)
        return ExtractedDocument(text=text, source_kind="ccda", warnings=warnings)
    if name.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp")):
        return _extract_image(data)
    # Unknown: best-effort decode.
    return ExtractedDocument(
        text=_decode(data), source_kind="text",
        warnings=[f"Unrecognized file type for {filename!r}; treated as plain text."],
    )


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="ignore")


def _strip_rtf(text: str) -> str:
    if "\\rtf" not in text[:64].lower():
        return text
    text = text.replace("\\par", "\n").replace("\\line", "\n").replace("\\tab", "\t")
    text = _RTF_HEX.sub(lambda m: bytes.fromhex(m.group(1)).decode("latin-1", "ignore"), text)
    text = text.replace("{", " ").replace("}", " ")
    text = _RTF_CONTROL.sub(" ", text)
    text = text.replace("\\", "\n")
    return re.sub(r"[ \t]+", " ", text).strip()


def _extract_pdf(data: bytes) -> ExtractedDocument:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ExtractedDocument(text="", source_kind="pdf",
                                 warnings=["PyMuPDF not installed; cannot read PDF."])
    warnings: list[str] = []
    ocr_used = False
    chunks: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            page_text = page.get_text("text").strip()
            if page_text:
                chunks.append(page_text)
            else:
                # No text layer -> likely a scanned page. Try OCR.
                ocr_text, ok, warn = _ocr_pdf_page(page)
                if ok:
                    chunks.append(ocr_text)
                    ocr_used = True
                elif warn:
                    warnings.append(warn)
    text = "\n\n".join(chunks)
    if not text and not warnings:
        warnings.append("No readable text was found in this PDF. If it is a scan, "
                        "paste the report text into the box below.")
    return ExtractedDocument(
        text=text, source_kind="pdf+ocr" if ocr_used else "pdf",
        warnings=warnings, ocr_used=ocr_used,
    )


def _ocr_pdf_page(page) -> tuple[str, bool, str | None]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError:
        return "", False, ("This looks like a scanned page — an image with no selectable "
                           "text. Paste the report text into the box below, or enable OCR "
                           "on the server (see README) to read scans automatically.")
    try:
        import fitz  # noqa: F401
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img), True, None
    except Exception as exc:  # pragma: no cover - depends on tesseract binary
        return "", False, f"OCR failed for a scanned page: {exc}"


def _extract_docx(data: bytes) -> ExtractedDocument:
    try:
        import docx  # python-docx
    except ImportError:
        return ExtractedDocument(text="", source_kind="docx",
                                 warnings=["python-docx not installed; cannot read DOCX."])
    document = docx.Document(io.BytesIO(data))
    paras = [p.text for p in document.paragraphs if p.text.strip()]
    # Include table cells (labs/meds are often tabular).
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return ExtractedDocument(text="\n".join(paras), source_kind="docx")


def _extract_image(data: bytes) -> ExtractedDocument:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError:
        return ExtractedDocument(
            text="", source_kind="image",
            warnings=["This image can't be read without OCR. Paste the report text into "
                      "the box below, or enable OCR on the server (see README)."],
        )
    try:
        img = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(img)
        return ExtractedDocument(text=text, source_kind="image", ocr_used=True)
    except Exception as exc:  # pragma: no cover
        return ExtractedDocument(text="", source_kind="image",
                                 warnings=[f"OCR failed: {exc}"])
